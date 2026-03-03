#!/usr/bin/env python3
"""
docx_builder.py — 템플릿 기반 DOCX 생성기

Usage:
  python docx_builder.py --content content.json [options]

Options:
  --content CONTENT_JSON    문서 구조 JSON 파일 (필수)
  --template TEMPLATE       빌트인 템플릿: report|memo|proposal|letter (기본: report)
  --custom-theme THEME_JSON 커스텀 테마 JSON 파일 (빌트인 위에 덮어씀)
  --output OUTPUT           출력 파일 경로 (기본: output.docx)

Content JSON 구조:
  {
    "title": "문서 제목",
    "meta": {
      "author": "홍길동",
      "date": "2024-03-01",
      "subject": "보고서 주제"
    },
    "sections": [
      {"type": "heading", "text": "1. 개요", "level": 1},
      {"type": "paragraph", "text": "본문 내용입니다."},
      {"type": "bullets", "items": ["항목 A", "항목 B"]},
      {"type": "numbers", "items": ["단계 1", "단계 2"]},
      {"type": "table", "headers": ["항목","값"], "rows": [["A","100"]]},
      {"type": "image", "path": "chart.png", "width_cm": 15},
      {"type": "page_break"}
    ]
  }

Custom theme JSON 구조 (모든 필드 선택적):
  {
    "heading1_color": "2E5984",
    "heading2_color": "4472C4",
    "heading3_color": "5B9BD5",
    "body_font":      "Malgun Gothic",
    "heading_font":   "Malgun Gothic",
    "font_size_body": 11,
    "font_size_h1":   16,
    "font_size_h2":   13,
    "font_size_h3":   11,
    "margin_top_cm":  2.5,
    "margin_bottom_cm": 2.5,
    "margin_left_cm": 3.0,
    "margin_right_cm": 3.0,
    "line_spacing":   1.15
  }
"""

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ── 빌트인 템플릿 ─────────────────────────────────────────────
BUILTIN_TEMPLATES: dict[str, dict] = {
    "report": {
        "heading1_color": "2E5984",
        "heading2_color": "4472C4",
        "heading3_color": "5B9BD5",
        "body_font":      "Malgun Gothic",
        "heading_font":   "Malgun Gothic",
        "font_size_body": 10.5,
        "font_size_h1":   16,
        "font_size_h2":   13,
        "font_size_h3":   11,
        "margin_top_cm":  2.5,
        "margin_bottom_cm": 2.5,
        "margin_left_cm": 3.0,
        "margin_right_cm": 3.0,
        "line_spacing":   1.15,
        "title_align":    "center",
    },
    "memo": {
        "heading1_color": "333333",
        "heading2_color": "555555",
        "heading3_color": "777777",
        "body_font":      "Malgun Gothic",
        "heading_font":   "Malgun Gothic",
        "font_size_body": 10,
        "font_size_h1":   13,
        "font_size_h2":   11,
        "font_size_h3":   10,
        "margin_top_cm":  2.0,
        "margin_bottom_cm": 2.0,
        "margin_left_cm": 2.5,
        "margin_right_cm": 2.5,
        "line_spacing":   1.0,
        "title_align":    "left",
    },
    "proposal": {
        "heading1_color": "1F497D",
        "heading2_color": "376092",
        "heading3_color": "4F81BD",
        "body_font":      "Malgun Gothic",
        "heading_font":   "Malgun Gothic",
        "font_size_body": 11,
        "font_size_h1":   18,
        "font_size_h2":   14,
        "font_size_h3":   12,
        "margin_top_cm":  3.0,
        "margin_bottom_cm": 3.0,
        "margin_left_cm": 3.5,
        "margin_right_cm": 3.5,
        "line_spacing":   1.5,
        "title_align":    "center",
    },
    "letter": {
        "heading1_color": "000000",
        "heading2_color": "333333",
        "heading3_color": "555555",
        "body_font":      "Batang",
        "heading_font":   "Batang",
        "font_size_body": 11,
        "font_size_h1":   13,
        "font_size_h2":   11,
        "font_size_h3":   10,
        "margin_top_cm":  3.0,
        "margin_bottom_cm": 3.0,
        "margin_left_cm": 4.0,
        "margin_right_cm": 4.0,
        "line_spacing":   1.5,
        "title_align":    "center",
    },
}


def hex_to_rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def load_template(template: str, custom_path: str | None) -> dict:
    tmpl = dict(BUILTIN_TEMPLATES.get(template, BUILTIN_TEMPLATES["report"]))
    if custom_path:
        overrides = json.loads(Path(custom_path).read_text(encoding="utf-8"))
        tmpl.update(overrides)
    return tmpl


def set_margins(doc: Document, t: dict) -> None:
    for section in doc.sections:
        section.top_margin    = Cm(t.get("margin_top_cm", 2.5))
        section.bottom_margin = Cm(t.get("margin_bottom_cm", 2.5))
        section.left_margin   = Cm(t.get("margin_left_cm", 3.0))
        section.right_margin  = Cm(t.get("margin_right_cm", 3.0))


def apply_paragraph_spacing(para, line_spacing: float) -> None:
    from docx.oxml.ns import qn
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), str(int(line_spacing * 240)))
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


# ── 콘텐츠 빌더 ──────────────────────────────────────────────
def add_document_title(doc: Document, title: str, meta: dict, t: dict) -> None:
    align_map = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT}
    align = align_map.get(t.get("title_align", "center"), WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(title)
    run.bold = True
    run.font.name = t["heading_font"]
    run.font.size = Pt(t["font_size_h1"] + 4)
    run.font.color.rgb = hex_to_rgb(t["heading1_color"])

    # 메타 정보
    if meta:
        meta_lines = []
        if "author" in meta:
            meta_lines.append(f"작성자: {meta['author']}")
        if "date" in meta:
            meta_lines.append(f"날짜: {meta['date']}")
        if "subject" in meta:
            meta_lines.append(f"주제: {meta['subject']}")
        if meta_lines:
            mp = doc.add_paragraph(" | ".join(meta_lines))
            mp.alignment = align
            for run in mp.runs:
                run.font.name = t["body_font"]
                run.font.size = Pt(t["font_size_body"])
                run.font.color.rgb = hex_to_rgb("888888")

    doc.add_paragraph()  # 빈 줄


def add_heading(doc: Document, text: str, level: int, t: dict) -> None:
    level = max(1, min(3, level))
    color_key = f"heading{level}_color"
    size_key   = f"font_size_h{level}"
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = t["heading_font"]
    run.font.size = Pt(t[size_key])
    run.font.color.rgb = hex_to_rgb(t[color_key])
    apply_paragraph_spacing(para, t.get("line_spacing", 1.15))


def add_paragraph(doc: Document, text: str, t: dict) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = t["body_font"]
    run.font.size = Pt(t["font_size_body"])
    apply_paragraph_spacing(para, t.get("line_spacing", 1.15))


def add_bullets(doc: Document, items: list[str], t: dict) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.runs[0] if p.runs else p.add_run(item)
        if not p.runs:
            run.text = item
        else:
            run.text = item
        run.font.name = t["body_font"]
        run.font.size = Pt(t["font_size_body"])


def add_numbers(doc: Document, items: list[str], t: dict) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.runs[0] if p.runs else p.add_run(item)
        run.text = item
        run.font.name = t["body_font"]
        run.font.size = Pt(t["font_size_body"])


def add_table(doc: Document, headers: list[str], rows: list[list[str]], t: dict) -> None:
    from docx.oxml import OxmlElement
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(h)
        run.text = h
        run.bold = True
        run.font.name = t["heading_font"]
        run.font.size = Pt(t["font_size_body"])

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data[: len(headers)]):
            cell = row_cells[i]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.name = t["body_font"]
                run.font.size = Pt(t["font_size_body"])

    doc.add_paragraph()  # 표 아래 여백


def add_image(doc: Document, path: str, width_cm: float, t: dict) -> None:
    if Path(path).exists():
        doc.add_picture(path, width=Cm(width_cm))
    else:
        add_paragraph(doc, f"[이미지 없음: {path}]", t)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


CONTENT_BUILDERS = {
    "heading":    lambda doc, section, t: add_heading(doc, section.get("text", ""), section.get("level", 1), t),
    "paragraph":  lambda doc, section, t: add_paragraph(doc, section.get("text", ""), t),
    "bullets":    lambda doc, section, t: add_bullets(doc, section.get("items", []), t),
    "numbers":    lambda doc, section, t: add_numbers(doc, section.get("items", []), t),
    "table":      lambda doc, section, t: add_table(doc, section.get("headers", []), section.get("rows", []), t),
    "image":      lambda doc, section, t: add_image(doc, section.get("path", ""), section.get("width_cm", 15), t),
    "page_break": lambda doc, section, t: add_page_break(doc),
}


# ── 메인 ─────────────────────────────────────────────────────
def main() -> None:
    parser = argparse.ArgumentParser(description="템플릿 기반 DOCX 생성기")
    parser.add_argument("--content",      required=True, help="문서 구조 JSON 파일")
    parser.add_argument("--template",     default="report",
                        choices=list(BUILTIN_TEMPLATES), help="빌트인 템플릿")
    parser.add_argument("--custom-theme", default=None, help="커스텀 테마 JSON 파일")
    parser.add_argument("--output",       default="output.docx", help="출력 파일")
    args = parser.parse_args()

    content: dict = json.loads(Path(args.content).read_text(encoding="utf-8"))
    t = load_template(args.template, args.custom_theme)

    doc = Document()
    set_margins(doc, t)

    # 제목 + 메타
    title = content.get("title", "")
    meta  = content.get("meta", {})
    if title:
        add_document_title(doc, title, meta, t)
        print(f"  [title] {title}")

    # 섹션 순차 처리
    for i, section in enumerate(content.get("sections", [])):
        section_type = section.get("type", "paragraph")
        builder = CONTENT_BUILDERS.get(section_type)
        if builder is None:
            print(f"  [skip] 알 수 없는 타입: {section_type} (섹션 #{i+1})", file=sys.stderr)
            continue
        builder(doc, section, t)
        label = section.get("text") or section.get("heading") or section_type
        print(f"  [{i+1}] {section_type}: {label}")

    doc.save(args.output)
    print(f"saved: {args.output}")


if __name__ == "__main__":
    main()
